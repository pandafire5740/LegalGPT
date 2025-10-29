"""Document processing service for extracting text from various file formats."""
import io
import re
from typing import Dict, Any, List, Optional
from pathlib import Path
import logging
from docx import Document
import PyPDF2
from datetime import datetime

from app.services.vector_store import VectorStore
from app.services.search_ingest import chunk_text as token_chunk_text


logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing legal documents and extracting content."""
    
    def __init__(self, vector_store: VectorStore):
        """Initialize document processor with vector store."""
        self.vector_store = vector_store
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from a Word document.
        
        Args:
            file_content: Binary content of the Word document
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            logger.info(f"Extracted text from Word document, length: {len(full_text)}")
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from Word document, error: {str(e)}")
            raise
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from a PDF document.
        
        Args:
            file_content: Binary content of the PDF document
            
        Returns:
            Extracted text as string
        """
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text.strip())
            
            full_text = "\n".join(text_parts)
            logger.info(f"Extracted text from PDF document, pages: {len(pdf_reader.pages)}, length: {len(full_text)}")
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF document, error: {str(e)}")
            raise
    
    def extract_text_from_text(self, file_content: bytes) -> str:
        """
        Extract text from a plain text file.
        
        Args:
            file_content: Binary content of the text file
            
        Returns:
            Extracted text as string
        """
        try:
            # Try UTF-8 first, then fallback to other encodings
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    text = file_content.decode('utf-8', errors='ignore')
            
            logger.info(f"Extracted text from text file, length: {len(text)}")
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract text from text file, error: {str(e)}")
            raise
    
    def extract_text(self, file_content: bytes, file_name: str) -> str:
        """
        Extract text from a document based on its file extension.
        
        Args:
            file_content: Binary content of the file
            file_name: Name of the file (used to determine type)
            
        Returns:
            Extracted text as string
        """
        file_extension = Path(file_name).suffix.lower()
        
        if file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_content)
        elif file_extension == '.pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_extension in ['.txt', '.rtf']:
            return self.extract_text_from_text(file_content)
        else:
            logger.warning(f"Unsupported file type, file_name=file_name, extension: {file_extension}")
            return ""
    
    def chunk_text(self, text: str, file_id: str = "", filename: str = "") -> List[str]:
        """
        Split text into overlapping chunks using token-based chunking.
        Uses the consolidated chunking function from search_ingest.
        
        Args:
            text: Text to chunk
            file_id: Optional file identifier for chunk metadata
            filename: Optional filename for chunk metadata
            
        Returns:
            List of text chunks
        """
        # Use token-based chunking from search_ingest
        chunk_dicts = token_chunk_text(text, file_id or "default", filename or "unknown")
        # Extract just the text content
        chunks = [chunk["text"] for chunk in chunk_dicts]
        logger.info(f"Chunked text, original_length: {len(text)}, chunks: {len(chunks)}")
        return chunks
    
    def extract_legal_entities(self, text: str) -> Dict[str, List[str]]:
        """
        Extract common legal entities and terms from text.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with different types of legal entities found
        """
        entities = {
            "parties": [],
            "dates": [],
            "amounts": [],
            "terms_conditions": [],
            "clauses": []
        }
        
        try:
            # Extract potential party names (capitalized words/phrases)
            party_pattern = r'\b[A-Z][A-Z\s&,\.]+(?:LLC|Inc\.|Corp\.|Ltd\.|Company|Co\.)\b'
            entities["parties"] = list(set(re.findall(party_pattern, text)))
            
            # Extract dates
            date_patterns = [
                r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
                r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{2,4}\b',
                r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4}\b'
            ]
            for pattern in date_patterns:
                entities["dates"].extend(re.findall(pattern, text, re.IGNORECASE))
            
            # Extract monetary amounts
            amount_pattern = r'\$[\d,]+\.?\d*'
            entities["amounts"] = list(set(re.findall(amount_pattern, text)))
            
            # Extract terms and conditions sections
            terms_pattern = r'(?:terms?\s+(?:and|&)\s+conditions?|t&c|terms?\s+of\s+(?:use|service))[^\n]*'
            entities["terms_conditions"] = re.findall(terms_pattern, text, re.IGNORECASE)
            
            # Extract common legal clauses
            clause_patterns = [
                r'(?:confidentiality|non-disclosure|nda)\s+(?:clause|agreement|provision)[^\n]*',
                r'(?:termination|cancellation)\s+(?:clause|provision)[^\n]*',
                r'(?:liability|indemnification)\s+(?:clause|provision)[^\n]*',
                r'(?:governing\s+law|jurisdiction)\s+(?:clause|provision)[^\n]*'
            ]
            for pattern in clause_patterns:
                entities["clauses"].extend(re.findall(pattern, text, re.IGNORECASE))
            
            logger.info(f"Extracted legal entities, parties: {len(entities['parties'])}, dates: {len(entities['dates'])}, amounts: {len(entities['amounts'])}, terms_conditions: {len(entities['terms_conditions'])}, clauses: {len(entities['clauses'])}")
            
            return entities
            
        except Exception as e:
            logger.error(f"Failed to extract legal entities, error: {str(e)}")
            return entities
    
    def process_document(self, file_content: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process a complete document: extract text, chunk it, and store in vector database.
        
        Args:
            file_content: Binary content of the file
            file_info: File metadata from Sharepoint
            
        Returns:
            Processing results dictionary
        """
        try:
            # Extract text
            text = self.extract_text(file_content, file_info["name"])
            if not text:
                logger.warning(f"No text extracted from document, file: {file_info['name']}")
                return {"status": "skipped", "reason": "No text extracted"}
            
            # Extract legal entities
            entities = self.extract_legal_entities(text)
            
            # Chunk text for vector storage
            chunks = self.chunk_text(text, file_id=file_info["name"], filename=file_info["name"])
            
            # Store in vector database
            doc_metadata = {
                "file_name": file_info["name"],
                "file_path": file_info["server_relative_url"],
                "time_modified": file_info["time_last_modified"],
                "author": file_info.get("author", "Unknown"),
                "file_size": file_info.get("length", 0),
                "entities": str(entities),  # Convert to string for ChromaDB compatibility
                "chunk_count": len(chunks)
            }
            
            # Store chunks with metadata
            chunk_ids = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = doc_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": i,
                    "chunk_text": chunk
                })
                
                chunk_id = f"{file_info['name']}_{i}"
                self.vector_store.add_document(chunk_id, chunk, chunk_metadata)
                chunk_ids.append(chunk_id)
            
            result = {
                "status": "success",
                "file_name": file_info["name"],
                "text_length": len(text),
                "chunks": len(chunks),
                "entities": entities,
                "chunk_ids": chunk_ids
            }
            
            logger.info(f"Successfully processed document, file: {file_info['name']}, text_length: {len(text)}, chunks: {len(chunks)}")
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to process document, file: {file_info['name']}, error: {str(e)}")
            return {"status": "error", "error": str(e)}
    
    def update_document(self, file_content: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Update an existing document in the vector database.
        
        Args:
            file_content: Binary content of the file
            file_info: File metadata from Sharepoint
            
        Returns:
            Update results dictionary
        """
        try:
            # First, remove existing chunks for this document
            self.vector_store.delete_document_chunks(file_info["name"])
            
            # Then process as new document
            return self.process_document(file_content, file_info)
            
        except Exception as e:
            logger.error(f"Failed to update document, file: {file_info['name']}, error: {str(e)}")
            return {"status": "error", "error": str(e)}
